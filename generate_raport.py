from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Creează documentul
doc = Document()

# Setează marginile standard
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# Titlu instituție - centrat
p = doc.add_paragraph("Ministerul Educaţiei și Cercetării al Republicii Moldova")
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph("Universitatea Tehnică a Moldovei")
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph("Facultatea Calculatoare, Informatică și Microelectronică")
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'

# Spații
doc.add_paragraph()
doc.add_paragraph()

# RAPORT
p = doc.add_paragraph("RAPORT")
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].font.size = Pt(14)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

p = doc.add_paragraph("Lucrarea de laborator nr.4")
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph('"Dezvoltarea Funcționalităților de Votare, Draft și Salvare"')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'

p = doc.add_paragraph("la Tehnologii WEB")
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'

# Student info
p = doc.add_paragraph()
run = p.add_run("A efectuat: \nst. gr. TI-243                                        Bularga Nichita")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()

# Profesor info
p = doc.add_paragraph()
run = p.add_run("A verificat:\nasistent univ.                                    Bunescu Mihai")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

# Spații
for _ in range(5):
    doc.add_paragraph()

# Chișinău
p = doc.add_paragraph("Chișinău 2026")
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

# Pagină nouă
doc.add_page_break()

# Tema lucrării
p = doc.add_paragraph()
run = p.add_run("Tema lucrării: ")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True
run = p.add_run("Implementarea funcționalităților de Votare (Vote), Draft și Salvare (SavedItem) în aplicația ForumApp Backend")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

# Scopul lucrării
p = doc.add_paragraph()
run = p.add_run("Scopul lucrării: ")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True
run = p.add_run("Dezvoltarea și integrarea componentelor funcționale pentru sistem de votare, gestionare draft-uri și salvare conținut, cu respectarea arhitecturii pe straturi și principiilor de securitate.")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

# Sarcinile lucrării
p = doc.add_paragraph()
run = p.add_run("Sarcinile lucrării:")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run.bold = True
p.paragraph_format.space_after = Pt(0)

tasks = [
    "Proiectarea entităților de domeniu pentru Vote, Draft și SavedItem cu relațiile corespunzătoare.",
    "Implementarea interfețelor de servicii pentru separarea logicii de business de controllere.",
    "Dezvoltarea serviciilor business cu validări complete și gestionarea excepțiilor.",
    "Crearea controllerelor RESTful cu endpoint-uri pentru operațiuni CRUD și interogări specifice.",
    "Implementarea mecanismelor de validare și autorizare la nivel de serviciu.",
    "Asigurarea integrității datelor prin actualizarea automată a contorilor (pentru voturi).",
    "Testarea funcționalităților și validarea comportamentului așteptat."
]

for task in tasks:
    p = doc.add_paragraph(f"• {task}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

# Introducere
p = doc.add_paragraph("Introducere")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True
p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("Prezentul raport documentează implementarea a trei funcționalități esențiale în cadrul aplicației ForumApp Backend: sistemul de votare (Vote), gestionarea draft-urilor (Draft) și salvarea conținutului (SavedItem). Aceste componente sunt dezvoltate folosind arhitectura pe straturi (layered architecture), respectând principiile SOLID și asigurând separarea clară a responsabilităților.")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

p = doc.add_paragraph("Fiecare funcționalitate este implementată prin patru straturi distincte: entități de domeniu (Domain), interfețe de servicii (Interfaces), servicii business (BusinessLayer) și controllere API (API Layer). Această abordare permite mentenanță ușoară, testabilitate crescută și extensibilitate optimă a aplicației.")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

doc.add_paragraph()

# === VOTE SECTION ===
p = doc.add_paragraph("1. Funcționalitatea Vote (Sistem de Votare)")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

doc.add_paragraph()

p = doc.add_paragraph("1.1. Entitatea de Domeniu")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

p = doc.add_paragraph("Entitatea VoteData reprezintă un vot acordat de un utilizator asupra unei postări sau unui comentariu. Structura permite asocierea unui vot fie cu un post, fie cu un comentariu, dar nu ambele simultan.")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

p = doc.add_paragraph("Proprietăți principale:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

properties = [
    "Id: Cheie primară cu generare automată (Identity), identificator unic al votului.",
    "Type: Tip de vot (VoteType enum) cu valorile DownVote (-1) și UpVote (1), utilizat pentru calculul scorului.",
    "PostId și Post: Relație opțională cu postarea votată (nullable).",
    "CommentId și Comment: Relație opțională cu comentariul votat (nullable).",
    "VotedAt: Data și ora când votul a fost acordat, cu valoare default DateTime.UtcNow.",
    "AuthorId și Author: Relație obligatorie cu utilizatorul care a acordat votul."
]

for prop in properties:
    p = doc.add_paragraph(f"• {prop}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("Logica de business asigură că fiecare vot este asociat fie cu un Post, fie cu un Comment, dar nu cu ambele. Enum-ul VoteType folosește valori numerice (-1 și 1) pentru facilitarea calculelor matematice ale scorului total.")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

doc.add_paragraph()

p = doc.add_paragraph("1.2. Modele DTO (Data Transfer Objects)")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

p = doc.add_paragraph("CreateVoteRequestDTO - utilizat pentru crearea unui vot nou:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

dto_create = [
    "Type: Tipul votului (UpVote sau DownVote), obligatoriu.",
    "PostId: ID-ul postării (opțional, nullable).",
    "CommentId: ID-ul comentariului (opțional, nullable)."
]

for item in dto_create:
    p = doc.add_paragraph(f"• {item}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("VoteResponseDTO - utilizat pentru returnarea datelor despre vot:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

dto_response = [
    "Id: Identificatorul votului.",
    "Type: Tipul votului.",
    "PostId și CommentId: Referințe către entitatea votată.",
    "VotedAt: Timestamp-ul votării.",
    "AuthorId și AuthorUserName: Informații despre autorul votului."
]

for item in dto_response:
    p = doc.add_paragraph(f"• {item}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("1.3. Interfața IVoteActions")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

p = doc.add_paragraph("Interfața definește contractul pentru operațiunile asupra voturilor, permițând injecția de dependențe și testabilitatea:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

methods = [
    "VoteAsync: Creează sau actualizează un vot existent.",
    "UpdateVoteAsync: Modifică tipul unui vot existent (schimbă între UpVote și DownVote).",
    "RemoveVoteAsync: Șterge un vot și actualizează contorul.",
    "GetVoteByIdAsync: Returnează detaliile unui vot specific.",
    "GetAllVotesAsync: Returnează toate voturile (pentru administrare).",
    "GetUserVoteOnPostAsync: Verifică votul unui utilizator pentru o postare specifică.",
    "GetUserVoteOnCommentAsync: Verifică votul unui utilizator pentru un comentariu specific."
]

for method in methods:
    p = doc.add_paragraph(f"• {method}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("1.4. Serviciul VoteService")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

p = doc.add_paragraph("VoteService implementează logica de business pentru gestionarea voturilor. Serviciul asigură integritatea datelor prin validări complexe și actualizare automată a contorilor.")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

p = doc.add_paragraph("Logica VoteAsync (crearea/actualizarea votului):")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

vote_logic = [
    "Validare inițială: verifică că există fie PostId, fie CommentId, dar nu ambele.",
    "Verificare existență: confirmă că postarea sau comentariul există în baza de date.",
    "Detectare vot existent: caută dacă utilizatorul a votat deja pe aceeași entitate.",
    "Dacă votul existent are același tip: returnează votul fără modificări.",
    "Dacă votul existent are tip diferit: calculează diferența (voteDifference = newValue - oldValue), actualizează votul și contorul.",
    "Dacă nu există vot: creează un vot nou, adaugă valoarea votului la contor.",
    "Actualizează contorul entității votate (Post sau Comment) folosind metoda UpdateVoteCounter.",
    "Salvează modificările și returnează VoteResponseDTO."
]

for logic in vote_logic:
    p = doc.add_paragraph(f"• {logic}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("Metoda UpdateVoteCounter:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

p = doc.add_paragraph("Această metodă helper actualizează contorul de voturi al postării sau comentariului. Primește ca parametru voteChange (diferența de voturi) și adaugă această valoare la proprietatea Votes a entității corespunzătoare. Acest mecanism asigură că scorul este mereu sincronizat cu voturile reale.")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

doc.add_paragraph()

p = doc.add_paragraph("Logica RemoveVoteAsync:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

remove_logic = [
    "Caută votul în baza de date după ID.",
    "Validează autorizarea: verifică că utilizatorul care șterge este autorul votului.",
    "Actualizează contorul folosind valoarea negativă a votului (pentru a-l scădea).",
    "Șterge votul din baza de date.",
    "Returnează ActionResponse cu status de succes sau eroare."
]

for logic in remove_logic:
    p = doc.add_paragraph(f"• {logic}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("1.5. VoteController")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

p = doc.add_paragraph("Controller-ul expune endpoint-uri RESTful pentru gestionarea voturilor, injectând serviciul IVoteActions prin Dependency Injection:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

endpoints = [
    "POST /api/vote - Creează sau actualizează un vot, primește CreateVoteRequestDTO și userId din query.",
    "PUT /api/vote/{voteId} - Actualizează tipul unui vot existent, cu validare de autorizare.",
    "DELETE /api/vote/{voteId} - Șterge un vot, doar dacă utilizatorul este autorul.",
    "GET /api/vote/{voteId} - Returnează detaliile unui vot specific.",
    "GET /api/vote - Returnează toate voturile (pentru administrare).",
    "GET /api/vote/post/{postId} - Returnează votul utilizatorului pe o postare specifică.",
    "GET /api/vote/comment/{commentId} - Returnează votul utilizatorului pe un comentariu specific."
]

for endpoint in endpoints:
    p = doc.add_paragraph(f"• {endpoint}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("Controller-ul efectuează validarea ModelState și returnează coduri HTTP corespunzătoare (200 OK, 400 BadRequest, 404 NotFound) în funcție de rezultatul operațiunii.")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

doc.add_paragraph()

# === DRAFT SECTION ===
p = doc.add_paragraph("2. Funcționalitatea Draft (Gestionare Draft-uri)")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

doc.add_paragraph()

p = doc.add_paragraph("2.1. Entitatea de Domeniu")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

p = doc.add_paragraph("Entitatea DraftData reprezintă o ciornă (draft) asociată unei postări, permițând utilizatorilor să salveze progresul lucrului la o postare înainte de publicare.")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

p = doc.add_paragraph("Proprietăți principale:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

properties = [
    "Id: Cheie primară cu generare automată (Identity).",
    "AuthorId și Author: Relație obligatorie cu utilizatorul care a creat draft-ul.",
    "PostId și Post: Relație obligatorie cu postarea la care se referă draft-ul.",
    "CreatedAt: Data și ora creării draft-ului, cu valoare default DateTime.UtcNow.",
    "LastModifiedAt: Data și ora ultimei modificări, actualizată automat la fiecare editare."
]

for prop in properties:
    p = doc.add_paragraph(f"• {prop}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("Draft-ul este privat și poate fi accesat doar de către autor. Entitatea asigură că un utilizator nu poate avea mai multe draft-uri pentru aceeași postare.")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

doc.add_paragraph()

p = doc.add_paragraph("2.2. Modele DTO")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

p = doc.add_paragraph("CreateDraftRequestDTO - utilizat pentru crearea unui draft nou:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

p = doc.add_paragraph("• PostId: ID-ul postării pentru care se creează draft-ul (obligatoriu).", style='List Bullet')
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("DraftResponseDTO - utilizat pentru returnarea datelor:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

dto_response = [
    "Id: Identificatorul draft-ului.",
    "AuthorId și AuthorUserName: Informații despre autorul draft-ului.",
    "PostId și PostTitle: Referință către postarea asociată.",
    "CreatedAt: Data creării.",
    "LastModifiedAt: Data ultimei modificări."
]

for item in dto_response:
    p = doc.add_paragraph(f"• {item}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("2.3. Interfața IDraftActions")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

p = doc.add_paragraph("Interfața definește operațiunile disponibile pentru draft-uri:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

methods = [
    "CreateDraftAsync: Creează un draft nou pentru o postare.",
    "UpdateDraftAsync: Actualizează draft-ul cu un PostId nou.",
    "GetDraftByIdAsync: Returnează detaliile unui draft specific (cu validare de autorizare).",
    "GetAllUserDraftsAsync: Returnează toate draft-urile unui utilizator.",
    "DeleteDraftAsync: Șterge un draft (doar autorul poate șterge)."
]

for method in methods:
    p = doc.add_paragraph(f"• {method}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("2.4. Serviciul DraftService")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

p = doc.add_paragraph("DraftService implementează logica de business pentru gestionarea draft-urilor cu validări stricte de securitate.")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

p = doc.add_paragraph("Logica CreateDraftAsync:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

create_logic = [
    "Verifică existența postării în baza de date.",
    "Caută un draft existent pentru aceeași postare și același autor.",
    "Dacă există deja: returnează draft-ul existent (prevenire duplicate).",
    "Dacă nu există: creează un draft nou cu timestamp-uri pentru CreatedAt și LastModifiedAt.",
    "Salvează draft-ul și încarcă relațiile (Author, Post) pentru response.",
    "Returnează DraftResponseDTO."
]

for logic in create_logic:
    p = doc.add_paragraph(f"• {logic}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("Logica UpdateDraftAsync:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

update_logic = [
    "Caută draft-ul după ID cu relațiile încărcate.",
    "Validare autorizare: verifică că utilizatorul este autorul draft-ului.",
    "Verifică existența noii postări.",
    "Actualizează PostId și LastModifiedAt.",
    "Salvează modificările și reîncarcă relația Post.",
    "Returnează DraftResponseDTO actualizat."
]

for logic in update_logic:
    p = doc.add_paragraph(f"• {logic}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("Toate metodele de citire (GetDraftByIdAsync, GetAllUserDraftsAsync) implementează validare de autorizare, asigurând că utilizatorii pot accesa doar propriile draft-uri. Metoda DeleteDraftAsync include aceeași validare înainte de ștergere.")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

doc.add_paragraph()

p = doc.add_paragraph("2.5. DraftController")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

p = doc.add_paragraph("Controller-ul expune endpoint-uri RESTful pentru gestionarea draft-urilor:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

endpoints = [
    "POST /api/draft - Creează un draft nou, primește CreateDraftRequestDTO și authorId din query.",
    "PUT /api/draft/{draftId} - Actualizează un draft existent (cu validare de autorizare).",
    "GET /api/draft/{draftId} - Returnează detaliile unui draft specific (doar pentru autor).",
    "GET /api/draft/user - Returnează toate draft-urile utilizatorului curent.",
    "DELETE /api/draft/{draftId} - Șterge un draft (doar autorul poate șterge)."
]

for endpoint in endpoints:
    p = doc.add_paragraph(f"• {endpoint}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("Controller-ul folosește try-catch pentru a gestiona InvalidOperationException (când postarea nu există) și returnează BadRequest cu mesaj descriptiv. Pentru operațiunile neautorizate sau resurse negăsite, returnează 404 NotFound.")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

doc.add_paragraph()

# === SAVEDITEM SECTION ===
p = doc.add_paragraph("3. Funcționalitatea SavedItem (Salvare Conținut)")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

doc.add_paragraph()

p = doc.add_paragraph("3.1. Entitatea de Domeniu")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

p = doc.add_paragraph("Entitatea SavedItemData reprezintă conținutul salvat de utilizatori pentru consultare ulterioară. Utilizatorii pot salva fie postări, fie comentarii.")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

p = doc.add_paragraph("Proprietăți principale:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

properties = [
    "Id: Cheie primară cu generare automată (Identity).",
    "AuthorId și Author: Relație obligatorie cu utilizatorul care a salvat conținutul.",
    "PostId și Post: Relație opțională cu postarea salvată (nullable).",
    "CommentId și Comment: Relație opțională cu comentariul salvat (nullable).",
    "CreatedAt: Data și ora salvării, cu valoare default DateTime.UtcNow."
]

for prop in properties:
    p = doc.add_paragraph(f"• {prop}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("Similar cu Vote, logica asigură că fiecare SavedItem este asociat fie cu un Post, fie cu un Comment, dar nu cu ambele. SavedItem-urile sunt private și vizibile doar pentru utilizatorul care le-a creat.")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

doc.add_paragraph()

p = doc.add_paragraph("3.2. Modele DTO")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

p = doc.add_paragraph("CreateSavedItemRequestDTO - utilizat pentru salvarea conținutului:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

dto_create = [
    "PostId: ID-ul postării de salvat (opțional, nullable).",
    "CommentId: ID-ul comentariului de salvat (opțional, nullable)."
]

for item in dto_create:
    p = doc.add_paragraph(f"• {item}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("SavedItemResponseDTO - utilizat pentru returnarea datelor:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

dto_response = [
    "Id: Identificatorul saved item-ului.",
    "AuthorId și AuthorUserName: Informații despre utilizatorul care a salvat.",
    "PostId și PostTitle: Referință către postarea salvată (dacă există).",
    "CommentId și CommentBody: Referință către comentariul salvat (dacă există).",
    "CreatedAt: Data salvării."
]

for item in dto_response:
    p = doc.add_paragraph(f"• {item}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("3.3. Interfața ISavedItemActions")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

p = doc.add_paragraph("Interfața definește operațiunile disponibile pentru conținut salvat:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

methods = [
    "SaveItemAsync: Salvează o postare sau un comentariu.",
    "RemoveSavedItemAsync: Șterge un saved item (doar proprietarul poate șterge).",
    "GetSavedItemByIdAsync: Returnează detaliile unui saved item specific (cu validare de autorizare).",
    "GetSavedItemsByUserAsync: Returnează toate saved items ale unui utilizator.",
    "GetUserSavedPostAsync: Verifică dacă un post specific este salvat de utilizator.",
    "GetUserSavedCommentAsync: Verifică dacă un comentariu specific este salvat de utilizator."
]

for method in methods:
    p = doc.add_paragraph(f"• {method}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("3.4. Serviciul SavedItemService")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

p = doc.add_paragraph("SavedItemService implementează logica de business pentru salvarea și gestionarea conținutului.")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

p = doc.add_paragraph("Logica SaveItemAsync:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

save_logic = [
    "Validare inițială: verifică că există fie PostId, fie CommentId, dar nu ambele.",
    "Verificare existență: confirmă că postarea sau comentariul există în baza de date.",
    "Detectare duplicat: caută dacă utilizatorul a salvat deja același conținut.",
    "Dacă conținutul este deja salvat: returnează saved item-ul existent (prevenire duplicate).",
    "Dacă nu există: creează un SavedItem nou cu timestamp CreatedAt.",
    "Salvează în baza de date și încarcă relațiile necesare (Author, Post sau Comment).",
    "Returnează SavedItemResponseDTO."
]

for logic in save_logic:
    p = doc.add_paragraph(f"• {logic}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("Logica RemoveSavedItemAsync:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

remove_logic = [
    "Caută saved item-ul după ID.",
    "Validare autorizare: verifică că utilizatorul este proprietarul saved item-ului.",
    "Șterge saved item-ul din baza de date.",
    "Returnează ActionResponse cu status de succes sau eroare."
]

for logic in remove_logic:
    p = doc.add_paragraph(f"• {logic}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("Metodele de interogare (GetUserSavedPostAsync, GetUserSavedCommentAsync) permit verificarea rapidă dacă un anumit conținut este salvat, util pentru UI (afișarea indicatorului \"salvat\").")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

doc.add_paragraph()

p = doc.add_paragraph("3.5. SavedItemController")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

p = doc.add_paragraph("Controller-ul expune endpoint-uri RESTful pentru gestionarea conținutului salvat:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

endpoints = [
    "POST /api/saveditem - Salvează o postare sau un comentariu.",
    "DELETE /api/saveditem/{savedItemId} - Șterge un saved item (doar proprietarul).",
    "GET /api/saveditem/{savedItemId} - Returnează detaliile unui saved item specific.",
    "GET /api/saveditem/user - Returnează toate saved items ale utilizatorului.",
    "GET /api/saveditem/post/{postId} - Verifică dacă un post este salvat de utilizator.",
    "GET /api/saveditem/comment/{commentId} - Verifică dacă un comentariu este salvat."
]

for endpoint in endpoints:
    p = doc.add_paragraph(f"• {endpoint}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("Controller-ul validează ModelState și returnează coduri HTTP semantice. Pentru resurse negăsite sau operațiuni neautorizate, returnează 404 NotFound cu mesaje descriptive.")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

doc.add_paragraph()

# === CONCLUZII ===
p = doc.add_paragraph("Concluzii")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

doc.add_paragraph()

p = doc.add_paragraph("În cadrul acestei lucrări au fost implementate cu succes trei funcționalități complexe pentru aplicația ForumApp Backend: sistemul de votare (Vote), gestionarea draft-urilor (Draft) și salvarea conținutului (SavedItem).")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

p = doc.add_paragraph("Implementarea respectă principiile arhitecturii pe straturi (layered architecture), asigurând separarea clară a responsabilităților între Domain, BusinessLayer și API Layer. Fiecare funcționalitate folosește interfețe pentru injectarea dependențelor, facilitând testabilitatea și mentenabilitatea codului.")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

p = doc.add_paragraph("Aspecte importante realizate:")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.runs[0].font.bold = True

conclusions = [
    "Validare robustă: Toate serviciile implementează validări complexe pentru asigurarea integrității datelor (verificare existență entități, prevenire duplicate, validare autorizare).",
    "Securitate: Implementarea autorizării la nivel de serviciu asigură că utilizatorii pot accesa și modifica doar propriile resurse (voturi, draft-uri, saved items).",
    "Integritate date: Sistemul de votare include mecanism automat de actualizare a contorilor, asigurând sincronizarea între voturi și scoruri.",
    "Prevenire duplicate: Logica de business previne crearea de duplicate pentru voturi și saved items, returnând resursele existente.",
    "Timestamp-uri: Toate entitățile folosesc DateTime.UtcNow pentru consistență temporală globală.",
    "Relații flexibile: Entitățile Vote și SavedItem suportă asocierea fie cu postări, fie cu comentarii, oferind flexibilitate maximă.",
    "API RESTful: Controller-ele expun endpoint-uri semantice cu coduri HTTP corespunzătoare, facilitând integrarea cu frontend.",
    "DTO pattern: Utilizarea modelelor DTO asigură expunerea controlată a datelor și separarea între modelul de domeniu și modelul de comunicare."
]

for conclusion in conclusions:
    p = doc.add_paragraph(f"• {conclusion}", style='List Bullet')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p = doc.add_paragraph("Aceste funcționalități constituie componente esențiale pentru orice aplicație de forum modernă, oferind utilizatorilor capacitatea de a interacționa cu conținutul (voturi), de a-și gestiona procesul de creare (draft-uri) și de a organiza conținutul de interes (saved items).")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

p = doc.add_paragraph("Arhitectura implementată permite extinderea ușoară cu noi funcționalități și integrarea cu sisteme de autentificare/autorizare (JWT) în etapele viitoare de dezvoltare.")
p.runs[0].font.size = Pt(12)
p.runs[0].font.name = 'Times New Roman'
p.paragraph_format.line_spacing = 1.5

# Salvează documentul
output_path = "C:\\Users\\user\\OneDrive\\Документы\\Tweb\\Raport Tweb - Vote Draft SavedItem.docx"
doc.save(output_path)

print(f"Raportul a fost generat cu succes la: {output_path}")
